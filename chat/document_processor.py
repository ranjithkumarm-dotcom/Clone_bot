"""
Document processing utilities for extracting text from various file formats
"""
import os
import mimetypes
from pathlib import Path
from io import BytesIO

def extract_text_from_file(file_input, file_type):
    """
    Extract text from various file formats

    Args:
        file_input: File object (Django UploadedFile) or file path
        file_type: MIME type or file extension

    Returns:
        Extracted text as string (empty string if extraction fails)
    """
    text = ""

    # Determine file extension from filename or file_input
    if hasattr(file_input, 'name'):
        filename = file_input.name
        file_ext = Path(filename).suffix.lower()
    elif isinstance(file_input, str):
        filename = file_input
        file_ext = Path(file_input).suffix.lower()
    else:
        file_ext = ''

    try:
        # PDF files
        if file_ext == '.pdf' or 'pdf' in file_type.lower():
            text = extract_from_pdf(file_input)

        # Word documents
        elif (file_ext in ['.docx', '.doc'] or
              'word' in file_type.lower() or
              'document' in file_type.lower()):
            text = extract_from_docx(file_input)

        # Excel files
        elif (file_ext in ['.xlsx', '.xls'] or
              'excel' in file_type.lower() or
              'spreadsheet' in file_type.lower()):
            text = extract_from_excel(file_input)

        # Text files
        elif file_ext in ['.txt', '.md', '.csv'] or 'text' in file_type.lower():
            text = extract_from_text(file_input)

        # Images (basic OCR would require additional libraries like pytesseract)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp'] or 'image' in file_type.lower():
            text = extract_from_image(file_input)

        else:
            # Try to read as text as fallback
            if hasattr(file_input, 'read'):
                try:
                    file_input.seek(0)
                    content = file_input.read()
                    if isinstance(content, bytes):
                        text = content.decode('utf-8', errors='ignore')
                    else:
                        text = content
                except Exception:
                    text = ""

    except Exception as e:
        # Return empty string on error, log the error instead
        import logging
        logging.error("Error processing file: %s", str(e))
        text = ""

    # Clean up the text - remove error messages that might have been returned
    if text and (text.startswith("Error") or text.startswith("Unable to")):
        text = ""

    return text.strip() if text else ""

def extract_from_pdf(file_input):
    """Extract text from PDF file using multiple methods for better compatibility.
    Prioritizes PyMuPDF (fitz) as it's most robust for various PDF types."""
    text = ""

    # Get file content as bytes
    if hasattr(file_input, 'read'):
        file_input.seek(0)
        file_content = file_input.read()
        file_input.seek(0)
    else:
        with open(file_input, 'rb') as f:
            file_content = f.read()

    # Method 1: Try PyMuPDF (fitz) FIRST - most robust, handles most PDF types
    try:
        import fitz  # PyMuPDF
        text = ""
        doc = fitz.open(stream=file_content, filetype="pdf")
        for page in doc:
            page_text = page.get_text()
            if page_text and page_text.strip():
                text += page_text + "\n"
        doc.close()
        if text.strip():
            return text.strip()
    except (ImportError, Exception):
        pass

    # Method 2: Try pdfplumber (better for complex PDFs and tables)
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
        if text.strip():
            return text.strip()
    except (ImportError, Exception):
        pass

    # Method 3: Try PyPDF2/pypdf as fallback
    try:
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n"
        except ImportError:
            # Try pypdf (newer version of PyPDF2)
            try:
                from pypdf import PdfReader
                pdf_reader = PdfReader(BytesIO(file_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
            except ImportError:
                pass

        if text.strip():
            return text.strip()
    except Exception:
        pass

    # If all methods fail, return empty string (not an error message)
    return ""

def extract_from_docx(file_input):
    """Extract text from DOCX file"""
    try:
        from docx import Document
        if hasattr(file_input, 'read'):
            file_input.seek(0)
            doc = Document(file_input)
        else:
            doc = Document(file_input)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception:
        return ""

def extract_from_excel(file_input):
    """Extract text from Excel file"""
    try:
        import openpyxl
        if hasattr(file_input, 'read'):
            file_input.seek(0)
            workbook = openpyxl.load_workbook(file_input)
        else:
            workbook = openpyxl.load_workbook(file_input)
        text = ""
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text += f"\n--- Sheet: {sheet_name} ---\n"
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    text += row_text + "\n"
        return text.strip()
    except Exception:
        return ""

def extract_from_text(file_input):
    """Extract text from plain text file"""
    try:
        if hasattr(file_input, 'read'):
            file_input.seek(0)
            content = file_input.read()
            if isinstance(content, bytes):
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        return content.decode(encoding).strip()
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                return ""
            return content.strip()
        else:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_input, 'r', encoding=encoding) as f:
                        return f.read().strip()
                except (UnicodeDecodeError, UnicodeError):
                    continue
            return ""
    except Exception:
        return ""

def extract_from_image(file_input):
    """Extract text from image (placeholder - would need OCR library)"""
    filename = file_input.name if hasattr(file_input, 'name') else str(file_input)
    return (
        f"Image file detected: {os.path.basename(filename)}. "
        f"OCR functionality requires additional setup (pytesseract). "
        f"Please provide text-based documents for now."
    )

def get_file_type(file_path):
    """Get MIME type of file"""
    mime_type, _ = mimetypes.guess_type(file_path)
    return mime_type or 'application/octet-stream'

